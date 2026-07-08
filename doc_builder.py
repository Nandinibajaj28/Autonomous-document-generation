import logging
import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

from models import AgentPlan, SectionContent

logger = logging.getLogger(__name__)

_OUTPUTS_DIR = Path(__file__).parent / "outputs"


def _slug(text: str) -> str:
    """Convert a title to a lowercase, underscore-separated filename-safe slug."""
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s-]+", "_", text)
    return text[:60]


def build_docx(plan: AgentPlan, sections: list[SectionContent]) -> str:
    """Assemble an AgentPlan and its generated sections into a .docx file.

    Writes a structured Word document with:
      - Heading 1 title, doc_type subtitle, and today's date
      - Heading 2 "Assumptions" block with one bullet per assumption
      - One Heading 2 + body paragraphs per SectionContent (split on blank lines)

    Saves to outputs/<slug(title)>_<YYYYMMDD>.docx and returns the path.

    Args:
        plan:     The AgentPlan supplying title, doc_type, and assumptions.
        sections: Ordered SectionContent list from executor.execute_plan().

    Returns:
        Absolute path string of the written .docx file.
    """
    _OUTPUTS_DIR.mkdir(exist_ok=True)

    doc = Document()

    doc.add_heading(plan.title, level=1)

    subtitle = doc.add_paragraph(plan.doc_type.title())
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(12)

    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))


    doc.add_heading("Assumptions", level=2)
    for assumption in plan.assumptions:
        doc.add_paragraph(assumption, style="List Bullet")

    for section in sections:
        doc.add_heading(section.heading, level=2)
        for paragraph_text in section.content.split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if paragraph_text:
                doc.add_paragraph(paragraph_text)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{_slug(plan.title)}_{timestamp}.docx"
    output_path = _OUTPUTS_DIR / filename
    doc.save(str(output_path))

    logger.info("Document saved: %s", output_path)
    return str(output_path)


if __name__ == "__main__":
    import sys
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        stream=sys.stdout,
    )

    from models import TaskStep

    fixture_plan = AgentPlan(
        doc_type="project plan",
        title="CRM Migration Project Plan",
        assumptions=[
            "The current CRM system and its functional requirements are well understood",
            "The target CRM system has been selected and its capabilities are known",
        ],
        steps=[
            TaskStep(
                step_id=1,
                title="Introduction",
                description="Overview of the CRM migration project.",
                section_heading="1. Introduction",
            ),
            TaskStep(
                step_id=2,
                title="Migration Strategy",
                description="Strategy for migrating data and system integration.",
                section_heading="2. Migration Strategy",
            ),
        ],
    )

    fixture_sections = [
        SectionContent(
            heading="1. Introduction",
            content=(
                "This project plan outlines the approach for migrating the organisation's "
                "CRM platform to a modern cloud-based solution. The migration aims to improve "
                "data quality, reduce operational overhead, and unlock advanced analytics capabilities.\n\n"
                "The scope covers data extraction, transformation, loading, integration testing, "
                "user acceptance testing, and a phased go-live rollout across all business units."
            ),
        ),
        SectionContent(
            heading="2. Migration Strategy",
            content=(
                "The migration will follow a phased approach, beginning with a pilot involving "
                "a representative subset of customer records. Each phase will be validated before "
                "the next begins, ensuring data integrity is preserved throughout.\n\n"
                "Automated ETL scripts will handle the bulk transfer, while a dedicated data-quality "
                "team will review edge cases and perform reconciliation checks against the source system.\n\n"
                "Rollback procedures are defined for each phase so that the original system can be "
                "reinstated within four hours if critical issues arise post-cutover."
            ),
        ),
    ]

    path = build_docx(fixture_plan, fixture_sections)
    print(f"\nDocument written to: {path}")
